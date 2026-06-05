from io import BytesIO
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet


def export_excel(meta, params, lph_rows, steps):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([meta]).to_excel(writer, sheet_name="Firmendaten", index=False)
        pd.DataFrame([params]).to_excel(writer, sheet_name="HOAI Parameter", index=False)
        pd.DataFrame(lph_rows).to_excel(writer, sheet_name="Leistungsphasen", index=False)
        pd.DataFrame(steps, columns=["Position", "Betrag €"]).to_excel(
            writer, sheet_name="Kalkulation", index=False
        )

    output.seek(0)
    return output


def export_word(meta, params, lph_rows, steps):
    output = BytesIO()
    doc = Document()

    doc.add_heading("ARCHiTool Honorarangebot", level=1)

    doc.add_heading("Firmendaten", level=2)
    for key, value in meta.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("HOAI Berechnungsgrundlagen", level=2)
    for key, value in params.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Leistungsphasen", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Leistungsphase"
    table.rows[0].cells[1].text = "Prozent"
    table.rows[0].cells[2].text = "Honorar netto €"

    for row in lph_rows:
        cells = table.add_row().cells
        cells[0].text = row["Leistungsphase"]
        cells[1].text = f"{row['Prozent']}%"
        cells[2].text = f"{row['Honorar netto €']:,.2f} €"

    doc.add_heading("Kademelige Kalkulation", level=2)
    for label, value in steps:
        doc.add_paragraph(f"{label}: {value:,.2f} €")

    doc.add_paragraph("")
    doc.add_paragraph("Dieses Dokument wurde automatisch mit ARCHiTool erstellt.")

    doc.save(output)
    output.seek(0)
    return output


def export_pdf(meta, params, lph_rows, steps):
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("ARCHiTool Honorarangebot", styles["Title"]))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Firmendaten", styles["Heading2"]))
    for key, value in meta.items():
        elements.append(Paragraph(f"<b>{key}:</b> {value}", styles["Normal"]))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph("HOAI Berechnungsgrundlagen", styles["Heading2"]))
    for key, value in params.items():
        elements.append(Paragraph(f"<b>{key}:</b> {value}", styles["Normal"]))

    elements.append(Spacer(1, 12))

    lph_table = [["Leistungsphase", "Prozent", "Honorar netto €"]]
    for row in lph_rows:
        lph_table.append([
            row["Leistungsphase"],
            f"{row['Prozent']}%",
            f"{row['Honorar netto €']:,.2f} €"
        ])

    table = Table(lph_table, colWidths=[280, 80, 120])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 18))

    calc_table = [["Position", "Betrag €"]]
    for label, value in steps:
        calc_table.append([label, f"{value:,.2f} €"])

    table2 = Table(calc_table, colWidths=[350, 130])
    table2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
    ]))
    elements.append(table2)

    doc.build(elements)
    output.seek(0)
    return output